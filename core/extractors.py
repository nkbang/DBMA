import re
"""
core/extractors.py — DBMA 텍스트 추출 엔진
지원 형식: pdf, txt, md, docx, epub, html/htm, rtf

추출 우선순위 (PDF):
  1순위: PyMuPDF (fitz)  — 폰트 매핑 정확도 최고, 한국어 CID 폰트 지원
  2순위: docling converter — OCR 포함 고품질 구조 분석
  3순위: pypdf            — 순수 Python fallback
  각 단계에서 추출 결과가 200자 미만이면 다음 단계로 자동 fallback

수정 이력:
  BUG-8  extract_text_from_pdf: converter 실패 시 예외 로그 기록
  BUG-9  PdfReader fallback 실패도 로그 기록
  BUG-10 extract_text_from_docx: 표(table) 셀 텍스트 추출 추가
  BUG-11 converter 미전달 시 docling 자동 빌드 옵션
  BUG-12 extract_text_from_epub: get_body_content() None 방어
  BUG-13 striprtf 미설치 시 ImportError graceful 처리
  NEW-1  PyMuPDF 1순위 추출 + 3단계 fallback 구조
  NEW-2  postprocess_pdf_text: PUA 제거, 한국어 어절 분리 복원, 고립 글리프 제거
  NEW-3  Tesseract OCR 추가 (he grc 언어 팩 지원)
  NEW-4  OCR 단어 내 공백 삽입 복원, 하이픈 줄바꿈 복원, 문단 경계 보호 (2026-06-29)
"""

import logging
import os
from typing import Dict, Optional

from bs4 import BeautifulSoup
from docx import Document
from ebooklib import epub, ITEM_DOCUMENT

logger = logging.getLogger(__name__)

# ─── 선택 의존성 안전 임포트 ─────────────────────────────
try:
    import fitz as _fitz
    _HAS_PYMUPDF = True
except ImportError:
    _fitz = None
    _HAS_PYMUPDF = False
    logger.warning("[EXTRACTORS] PyMuPDF 없음 — pip install pymupdf 권장")

try:
    from pypdf import PdfReader
    _HAS_PYPDF = True
except Exception:
    PdfReader = None
    _HAS_PYPDF = False
    logger.warning("[EXTRACTORS] pypdf 없음 — PDF pypdf fallback 비활성")

try:
    from striprtf.striprtf import rtf_to_text as _rtf_to_text
    _HAS_STRIPRTF = True
except ImportError:
    _HAS_STRIPRTF = False
    logger.warning("[EXTRACTORS] striprtf 없음 — RTF 추출 비활성")

try:
    import pytesseract
    _HAS_PYTESSERACT = True
except ImportError:
    _HAS_PYTESSERACT = False
    logger.warning("[EXTRACTORS] pytesseract 없음 — OCR 비활성")

try:
    from pdf2image import convert_from_path
    _HAS_PDF2IMAGE = True
except ImportError:
    _HAS_PDF2IMAGE = False
    logger.warning("[EXTRACTORS] pdf2image 없음 — OCR 비활성")


# ─────────────────────────────────────────────────────────
# 공통 텍스트 파일 읽기 (인코딩 폴백)
# ─────────────────────────────────────────────────────────
def read_text_file(path: str) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp949", "euc-kr"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, LookupError):
            continue
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


# ─────────────────────────────────────────────────────────
# 형식별 추출 함수
# ─────────────────────────────────────────────────────────
def extract_text_from_txt(path: str) -> str:
    return read_text_file(path).strip()


def extract_text_from_md(path: str) -> str:
    return read_text_file(path).strip()


def extract_text_from_docx(path: str) -> str:
    """
    BUG-10 fix: doc.paragraphs 뿐 아니라 표(table) 셀 텍스트도 추출.
    구조: 문단 → 빈줄 구분, 표 → 각 행을 탭으로 연결 후 빈줄 구분
    """
    doc   = Document(path)
    parts = []

    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))

    return "\n\n".join(parts).strip()


def extract_text_from_epub(path: str) -> str:
    """BUG-12 fix: get_body_content() None 반환 방어."""
    book  = epub.read_epub(path)
    texts = []
    for item in book.get_items():
        if item.get_type() != ITEM_DOCUMENT:
            continue
        try:
            body = item.get_body_content()
            if not body:
                continue
            soup = BeautifulSoup(body, "html.parser")
            text = soup.get_text("\n", strip=True)
            if text:
                texts.append(text)
        except Exception as e:
            logger.warning(f"[EPUB] 아이템 파싱 실패 {item.get_name()}: {e}")
    return "\n\n".join(texts).strip()


def extract_text_from_html(path: str) -> str:
    html = read_text_file(path)
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript", "head"]):
        tag.decompose()
    return soup.get_text("\n", strip=True).strip()


def extract_text_from_rtf(path: str) -> str:
    """BUG-13 fix: striprtf 미설치 시 명확한 예외 발생."""
    if not _HAS_STRIPRTF:
        raise ImportError(
            "RTF 추출에는 striprtf 가 필요합니다: pip install striprtf"
        )
    raw = read_text_file(path)
    return _rtf_to_text(raw).strip()


# ─────────────────────────────────────────────────────────
# PDF 텍스트 후처리 (NEW-4: 전면 개선 2026-06-29)
# ─────────────────────────────────────────────────────────
_RE_PUA_F        = re.compile(r'[\uE000-\uF8FF\U000F0000-\U000FFFFF]')
_RE_CTRL_F       = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]')
_RE_NBSP_F       = re.compile(r'[\xa0\u00ad\u200b\u200c\u200d\ufeff]')
_RE_ISOLATED     = re.compile(r'^\s*(?:[lIifl·•※◆▶→←↑↓]|ff|fi|fl)\s*$', re.MULTILINE)
_RE_MULTIBLANK_F = re.compile(r'\n{3,}')

# 하이픈 줄바꿈: "fulfil-\nment" → "fulfillment"
_RE_EN_HYPHEN_BREAK = re.compile(r'([A-Za-z])-\n([a-z])')

# 문단 경계 보호: 마침표/!/? 뒤 대문자 시작은 새 문단(\n\n) 유지
_RE_PARA_BOUNDARY = re.compile(r'([.!?])\n([A-Z\uAC00-\uD7A3])')

# 영문 soft-break: 줄 끝 영문자 + \n + 영문자 시작 → 공백으로 연결
_RE_EN_SOFTBREAK = re.compile(r'([a-zA-Z])\n([a-zA-Z])')

# 한국어 조사/어미 줄바꿈 복원
_KO_JOSA_PATTERN = re.compile(
    r'([\uAC00-\uD7A3])\n'
    r'(이|가|을|를|은|는|의|에서|에게|에|으로|로|와|과|도|만|까지|보다|처럼|부터|한테|께|서|고|며|면|야|랑|이랑|란|이란|만큼|대로|같이|님|적)'
    r'(?=[\uAC00-\uD7A3\s.,!?;:\n]|$)'
)


def _fix_ocr_word_splits(text: str) -> str:
    """
    OCR 단어 내 공백 삽입 복원.
    예)  "T he"          → "The"
         "m em bers"     → "members"
         "com m entary"  → "commentary"
         "o f"           → "of"
         "w ritten"      → "written"

    보호: 독립 관사 'a', 'A', 'I' 는 단어 경계 유지
    """
    t = text
    # 1단계: 단일 대문자 + 공백 + 소문자 2자 이상 ("T he" → "The")
    t = re.sub(r'\b([A-Z]) ([a-z]{2,})\b', r'\1\2', t)
    # 2단계: 단일 소문자(a·I 제외) + 공백 + 소문자 1-2자 ("o f" → "of")
    t = re.sub(r'\b([b-hj-z]) ([a-z]{1,2})\b', r'\1\2', t)
    # 3단계: 중간 단음절 분리 반복 제거 ("com m ent" → "comment")
    for _ in range(5):
        prev = t
        t = re.sub(r'(?<=[a-z]) ([a-z]{1,3})(?= [a-z])', r'\1', t)
        if t == prev:
            break
    return t


def postprocess_pdf_text(text: str) -> str:
    """
    PDF / OCR 추출 텍스트 최종 후처리.

    처리 순서:
    1. Unicode PUA 제거 (폰트 매핑 실패 잔재)
    2. 제어문자, 비표준 공백 정규화
    3. 고립 글리프 라인 제거 (l, I, ff 등)
    4. 하이픈 줄바꿈 복원 ("fulfil-\nment" → "fulfillment")
    5. 문단 경계 보호 후 영문 soft-break 공백 연결
    6. OCR 단어 내 공백 삽입 복원 ("T he" → "The")
    7. 한국어 조사/어미 줄바꿈 복원
    8. 연속 빈줄 정리
    """
    t = _RE_PUA_F.sub('', text)
    t = _RE_CTRL_F.sub('', t)
    t = _RE_NBSP_F.sub(' ', t)
    t = t.replace('\x0c', '\n')
    t = _RE_ISOLATED.sub('', t)

    # 4. 하이픈 줄바꿈 복원
    t = _RE_EN_HYPHEN_BREAK.sub(r'\1\2', t)

    # 5. 문단 경계 보호 후 soft-break 병합
    t = _RE_PARA_BOUNDARY.sub(r'\1\n\n\2', t)
    t = _RE_EN_SOFTBREAK.sub(r'\1 \2', t)

    # 6. OCR 단어 내 공백 복원
    t = _fix_ocr_word_splits(t)

    # 7. 한국어 조사/어미 복원 (수렴 감지)
    for _ in range(5):
        prev = t
        t = _KO_JOSA_PATTERN.sub(lambda m: m.group(1) + m.group(2), t)
        if t == prev:
            break

    # 8. 연속 빈줄 정리
    t = _RE_MULTIBLANK_F.sub('\n\n', t)
    return t.strip()


# ─────────────────────────────────────────────────────────
# PDF 추출 — 3단계 fallback
#   1순위: PyMuPDF  (fitz)       — 폰트 정확도 최고
#   2순위: docling  converter    — OCR 포함 고품질
#   3순위: pypdf                 — 순수 Python fallback
# ─────────────────────────────────────────────────────────
_MIN_TEXT_LEN = 200


def _extract_via_pymupdf(path: str) -> str:
    """
    PyMuPDF(fitz) 로 텍스트 추출.
    get_text("blocks") 모드: 단락 단위 추출로 줄바꿈 어절 분리 최소화.
    이미지 블록(b[6]==1)은 제외.
    """
    doc   = _fitz.open(path)
    pages = []
    for page in doc:
        blocks    = page.get_text("blocks")
        page_text = "\n".join(
            b[4].strip()
            for b in blocks
            if b[6] == 0 and b[4].strip()
        )
        if page_text.strip():
            pages.append(page_text)
    doc.close()
    return "\n\n".join(pages).strip()


def _extract_via_docling(path: str, converter) -> str:
    """docling DocumentConverter 로 텍스트 추출."""
    result = converter.convert(path)
    doc    = result.document
    if hasattr(doc, "export_to_markdown"):
        return doc.export_to_markdown() or ""
    if hasattr(doc, "export_to_text"):
        return doc.export_to_text() or ""
    return str(doc) or ""


def _extract_via_pypdf(path: str) -> str:
    """pypdf PdfReader 로 텍스트 추출."""
    reader = PdfReader(path)
    pages  = []
    for page in reader.pages:
        t = page.extract_text() or ""
        if t.strip():
            pages.append(t.strip())
    return "\n\n".join(pages).strip()


def _detect_ocr_flag(converter) -> bool:
    """docling converter 의 OCR 설정 여부 감지."""
    try:
        opts = getattr(converter, "format_options", None)
        if opts:
            pdf_opts = opts.get("pdf")
            if pdf_opts and getattr(pdf_opts, "pipeline_options", None):
                return bool(getattr(pdf_opts.pipeline_options, "do_ocr", False))
    except Exception:
        pass
    return False


def _extract_via_ocr(path: str) -> str:
    """Tesseract OCR 를 사용하여 PDF 이미지에서 텍스트 추출."""
    if not (_HAS_PYTESSERACT and _HAS_PDF2IMAGE):
        logger.warning("[PDF] OCR 의존성 없음 (pytesseract 또는 pdf2image)")
        return ""
    try:
        images = convert_from_path(path, dpi=300)
        full_text = []
        for img in images:
            txt = pytesseract.image_to_string(
                img,
                lang="kor+eng+heb+grc",
                config="--oem 3 --psm 6",
            )
            if txt.strip():
                full_text.append(txt.strip())
        return "\n\n".join(full_text).strip()
    except Exception as e:
        logger.warning(f"[PDF] OCR 추출 실패 ({os.path.basename(path)}): {e}")
        return ""


# ─────────────────────────────────────────────────────────
# 메인 PDF 추출 함수
# ─────────────────────────────────────────────────────────
def extract_text_from_pdf(
    path: str,
    converter=None,
    use_ocr: bool = False,
) -> str:
    """
    PDF 텍스트 추출 (3단계 fallback).

    Args:
        path:      PDF 파일 경로
        converter: docling DocumentConverter (선택, OCR 포함 고품질)
        use_ocr:   True 이면 Tesseract OCR 3순위로 활성화

    Returns:
        추출된 텍스트 (postprocess_pdf_text 적용 완료)
    """
    text = ""

    # ── 1순위: PyMuPDF ──────────────────────────────────
    if _HAS_PYMUPDF:
        try:
            text = _extract_via_pymupdf(path)
            if len(text) >= _MIN_TEXT_LEN:
                logger.info(
                    f"[PDF] PyMuPDF 추출 성공: {os.path.basename(path)} "
                    f"({len(text):,}자)"
                )
        except Exception as e:
            logger.warning(f"[PDF] PyMuPDF 실패 ({os.path.basename(path)}): {e}")
            text = ""

    # ── 2순위: docling ───────────────────────────────────
    if len(text) < _MIN_TEXT_LEN and converter is not None:
        try:
            text = _extract_via_docling(path, converter)
            if len(text) >= _MIN_TEXT_LEN:
                logger.info(
                    f"[PDF] docling 추출 성공: {os.path.basename(path)} "
                    f"({len(text):,}자)"
                )
        except Exception as e:
            logger.warning(f"[PDF] docling 실패 ({os.path.basename(path)}): {e}")
            text = ""

    # ── 3순위: OCR (Tesseract) ───────────────────────────
    if len(text) < _MIN_TEXT_LEN and use_ocr:
        try:
            text = _extract_via_ocr(path)
            if len(text) >= _MIN_TEXT_LEN:
                logger.info(
                    f"[PDF] OCR 추출 성공: {os.path.basename(path)} "
                    f"({len(text):,}자)"
                )
        except Exception as e:
            logger.warning(f"[PDF] OCR 실패 ({os.path.basename(path)}): {e}")
            text = ""

    # ── 4순위: pypdf ─────────────────────────────────────
    if len(text) < _MIN_TEXT_LEN and _HAS_PYPDF:
        try:
            text = _extract_via_pypdf(path)
            if text:
                logger.info(
                    f"[PDF] pypdf fallback 성공: {os.path.basename(path)} "
                    f"({len(text):,}자)"
                )
        except Exception as e:
            logger.warning(f"[PDF] pypdf 실패 ({os.path.basename(path)}): {e}")
            text = ""

    if not text.strip():
        raise ValueError(
            f"PDF 텍스트 추출 완전 실패: {os.path.basename(path)}\n"
            "PyMuPDF / docling / OCR / pypdf 모두 실패했습니다.\n"
            "스캔본이면 docling OCR 설정 또는 poppler 설치를 확인하세요."
        )

    text = postprocess_pdf_text(text)
    logger.info(
        f"[PDF] 후처리 완료: {os.path.basename(path)} → {len(text):,}자"
    )
    return text


# ─────────────────────────────────────────────────────────
# 통합 추출 라우터
# ─────────────────────────────────────────────────────────
def extract_text_from_file(
    path: str,
    converter=None,
    use_ocr: bool = False,
) -> dict:
    """
    파일 확장자에 따라 적절한 추출 함수를 호출한다.

    지원 형식: .pdf .txt .md .docx .epub .html .htm .rtf

    Args:
        path:      대상 파일 경로
        converter: docling DocumentConverter (PDF OCR 용, 선택)
        use_ocr:   True 이면 Tesseract OCR 활성화

    Returns:
        dict: {"text": str, "is_ocr": bool, "source_type": str}
    """
    ext = os.path.splitext(path)[1].lower()

    text = ""
    is_ocr = False

    if ext == ".pdf":
        text = extract_text_from_pdf(path, converter=converter, use_ocr=use_ocr)
        is_ocr = use_ocr or _detect_ocr_flag(converter) if converter else use_ocr
    elif ext == ".txt":
        text = extract_text_from_txt(path)
    elif ext == ".md":
        text = extract_text_from_md(path)
    elif ext == ".docx":
        text = extract_text_from_docx(path)
    elif ext == ".epub":
        text = extract_text_from_epub(path)
    elif ext in (".html", ".htm"):
        text = extract_text_from_html(path)
    elif ext == ".rtf":
        text = extract_text_from_rtf(path)
    else:
        raise ValueError(
            f"지원하지 않는 파일 형식: '{ext}'\n"
            f"지원 형식: pdf, txt, md, docx, epub, html, htm, rtf"
        )

    return {"text": text, "is_ocr": is_ocr, "source_type": ext.lstrip(".")}
