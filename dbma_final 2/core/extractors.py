"""
core/extractors.py — DBMA 텍스트 추출 엔진
지원 형식: pdf, txt, md, docx, epub, html/htm, rtf

디버그 수정 사항:
  BUG-8  extract_text_from_pdf: converter 실패 시 예외를 로그에 기록 (무음 삼킴 제거)
  BUG-9  PdfReader fallback 실패도 로그에 기록
  BUG-10 extract_text_from_docx: 표(table) 내 텍스트 누락 → 표 셀도 추출
  BUG-11 converter 미전달 시 docling 자동 빌드 옵션 추가 (dbma.py 쪽에서 주입 가능)
  BUG-12 extract_text_from_epub: get_body_content() None 방어
  BUG-13 striprtf 미설치 시 ImportError graceful 처리
"""

import logging
import os
from typing import Dict, Optional

from bs4 import BeautifulSoup
from docx import Document
from ebooklib import epub, ITEM_DOCUMENT

logger = logging.getLogger(__name__)

# ─── 선택 의존성 안전 임포트 ─────────────────────────────
try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None
    logger.warning("[EXTRACTORS] pypdf 없음 — PDF pypdf fallback 비활성")

try:
    from striprtf.striprtf import rtf_to_text as _rtf_to_text
    _HAS_STRIPRTF = True
except ImportError:
    _HAS_STRIPRTF = False
    logger.warning("[EXTRACTORS] striprtf 없음 — RTF 추출 비활성")


# ─────────────────────────────────────────────────────────
# 공통 텍스트 파일 읽기 (인코딩 폴백)
# ─────────────────────────────────────────────────────────
def read_text_file(path: str) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp949", "euc-kr"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, LookupError):
            continue
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


# ─────────────────────────────────────────────────────────
# 형식별 추출 함수
# ─────────────────────────────────────────────────────────
def extract_text_from_txt(path: str) -> str:
    return read_text_file(path).strip()


def extract_text_from_md(path: str) -> str:
    return read_text_file(path).strip()


def extract_text_from_docx(path: str) -> str:
    """
    BUG-10 fix: doc.paragraphs 뿐 아니라 표(table) 셀 텍스트도 추출.
    구조: 문단 → 빈줄 구분, 표 → 각 행을 탭으로 연결 후 빈줄 구분
    """
    doc   = Document(path)
    parts = []

    # 1) 일반 문단
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)

    # 2) 표 셀
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))

    return "\n\n".join(parts).strip()


def extract_text_from_epub(path: str) -> str:
    """BUG-12 fix: get_body_content() None 반환 방어."""
    book  = epub.read_epub(path)
    texts = []
    for item in book.get_items():
        if item.get_type() != ITEM_DOCUMENT:
            continue
        try:
            body = item.get_body_content()
            if not body:          # None 또는 빈 bytes
                continue
            soup = BeautifulSoup(body, "html.parser")
            text = soup.get_text("\n", strip=True)
            if text:
                texts.append(text)
        except Exception as e:
            logger.warning(f"[EPUB] 아이템 파싱 실패 {item.get_name()}: {e}")
    return "\n\n".join(texts).strip()


def extract_text_from_html(path: str) -> str:
    html = read_text_file(path)
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript", "head"]):
        tag.decompose()
    return soup.get_text("\n", strip=True).strip()


def extract_text_from_rtf(path: str) -> str:
    """BUG-13 fix: striprtf 미설치 시 명확한 예외 발생."""
    if not _HAS_STRIPRTF:
        raise ImportError(
            "RTF 추출에는 striprtf 가 필요합니다: pip install striprtf"
        )
    raw = read_text_file(path)
    return _rtf_to_text(raw).strip()


# ─────────────────────────────────────────────────────────
# PDF 추출 (docling converter → pypdf fallback)
# BUG-8, BUG-9: 실패 원인을 로그에 기록하고 계속 시도
# ─────────────────────────────────────────────────────────
def extract_text_from_pdf(path: str, converter=None) -> Dict:
    """
    converter: docling DocumentConverter 인스턴스 (선택).
               None 이면 pypdf 만 사용.
    반환: {"text": str, "file_type": "pdf", "is_ocr": bool}
    """
    text   = ""
    is_ocr = False

    # ── docling converter 경로 ──────────────────────────
    if converter is not None:
        try:
            result = converter.convert(path)
            doc    = result.document

            if hasattr(doc, "export_to_markdown"):
                text = doc.export_to_markdown() or ""
            elif hasattr(doc, "export_to_text"):
                text = doc.export_to_text() or ""
            else:
                text = str(doc) or ""

            # OCR 여부 감지
            try:
                conv_opts = getattr(converter, "format_options", None)
                if conv_opts:
                    pdf_opts = conv_opts.get("pdf")
                    if pdf_opts and getattr(pdf_opts, "pipeline_options", None):
                        is_ocr = bool(getattr(pdf_opts.pipeline_options, "do_ocr", False))
            except Exception:
                pass  # OCR 감지 실패는 치명적이지 않음

        except Exception as e:
            # BUG-8 fix: 무음 삼킴 → 경고 로그
            logger.warning(f"[PDF] docling converter 실패 ({path}): {e} — pypdf fallback 시도")
            text = ""

    # ── pypdf fallback ──────────────────────────────────
    if not text.strip():
        if PdfReader is None:
            logger.error("[PDF] pypdf 미설치. pip install pypdf")
        else:
            try:
                reader = PdfReader(path)
                pages  = []
                for page in reader.pages:
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        pages.append(page_text.strip())
                text = "\n\n".join(pages).strip()
            except Exception as e:
                # BUG-9 fix: 무음 삼킴 → 경고 로그
                logger.warning(f"[PDF] pypdf 실패 ({path}): {e}")
                text = ""

    if not text.strip():
        raise ValueError(
            f"PDF 텍스트 추출 실패: {os.path.basename(path)}\n"
            "원인: docling converter 와 pypdf 모두 실패했습니다.\n"
            "조치: 파일이 스캔본이면 docling OCR 설정 확인 또는 poppler 설치 확인."
        )

    return {"text": text, "file_type": "pdf", "is_ocr": is_ocr}


# ─────────────────────────────────────────────────────────
# 통합 진입점
# BUG-11 참고: converter 를 외부에서 주입해야 PDF OCR 작동
# ─────────────────────────────────────────────────────────
def extract_text_from_file(
    src_path: str,
    file_type: str = "",
    converter=None,
) -> Dict:
    """
    파일 경로와 형식을 받아 텍스트를 추출한다.

    Args:
        src_path:  절대 경로
        file_type: 확장자 힌트 (없으면 파일명에서 자동 감지)
        converter: docling DocumentConverter (PDF OCR 용, 선택)

    Returns:
        {"text": str, "file_type": str, "is_ocr": bool}
    """
    ext = (
        file_type
        or os.path.splitext(src_path)[1]
    ).lower().lstrip(".").strip()

    if ext == "pdf":
        return extract_text_from_pdf(src_path, converter)

    if ext == "txt":
        return {"text": extract_text_from_txt(src_path), "file_type": "txt", "is_ocr": False}

    if ext == "md":
        return {"text": extract_text_from_md(src_path),  "file_type": "md",  "is_ocr": False}

    if ext == "docx":
        return {"text": extract_text_from_docx(src_path), "file_type": "docx", "is_ocr": False}

    if ext == "epub":
        return {"text": extract_text_from_epub(src_path), "file_type": "epub", "is_ocr": False}

    if ext in ("html", "htm"):
        return {"text": extract_text_from_html(src_path), "file_type": "html", "is_ocr": False}

    if ext == "rtf":
        return {"text": extract_text_from_rtf(src_path), "file_type": "rtf", "is_ocr": False}

    raise ValueError(
        f"지원하지 않는 파일 형식: '{ext}'\n"
        f"지원 형식: pdf, txt, md, docx, epub, html, htm, rtf"
    )
